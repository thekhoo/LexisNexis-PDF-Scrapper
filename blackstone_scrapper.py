from curses.ascii import isupper
import os
import json
import PyPDF2
import logging
import regex as re
from sys import stdout

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK

from typing import List

class Logger:
    
    def __init__(self) -> None:
        self.log = logging.getLogger()
        self.log.setLevel(logging.INFO)

        handler = logging.StreamHandler(stdout)
        handler.setLevel(logging.INFO)
        self.log.addHandler(handler)

    def logInfo(self,text:str) -> None:
        self.log.info(text)

class ProcedurePDF:

    SECTION_MAIN_HEADING = "section_heading"
    SECTION_SUB_HEADING = "section_subheading"
    SECTION_TEXT = "section_text"

    """
    Class to handle a SINGLE (.pdf) file from the Procedure (Part D) from Blackstone's Criminal Practice 2022 from Lexis Library
    """

    def __init__(self,filename:str) -> None:
        self.filename = filename
        self.pdf_text = self._getPDFText()
        self.pdf_dict = self._getPDFDict()

    def getSections(self,sections:List[int]) -> dict:
        """
        Returns a dictionary of the sections that are entered as a list of integers in this function.
        """
        all_section_dict = {}
        
        # Loop through all the section numbers provided
        for section in sections:
            key = f"{self.filename}.{section}"                      # Get the Section Heading
            section_data = self.pdf_dict[key]   # Use the Section heading as the key and remove newline characters

            all_section_dict.update({key: section_data})    # Update the main dictionary with the section dictionary
        
        return all_section_dict

    # Get's the PDF file and converts the text into a string
    def _getPDFText(self) -> str:
        """
        Get all the text from the PDF file in a single string.
        """
        # Open the File
        filepath = os.path.join('data',f'{self.filename}.pdf')

        try:
            file = open(filepath,'rb')
        except FileNotFoundError:
            print('We cannot find the file that you are looking for. Please try again.')
            return None
        else:
            # Create Reader Object
            reader = PyPDF2.PdfFileReader(file)

            # Get the Text
            text = ''

            for i in range(reader.numPages):
                page = reader.getPage(i)                # Get the Current Page
                current_text = page.extract_text()      # Extract the text from the current page
                text += current_text                    # Append to master 'text' variable
                text += '\n'                            # Add New Line character to the end of every page

            file.close()
            
            return text

    # Converts a string of text into a dictionary based on sections (i.e. D5.4)
    def _getPDFDict(self) -> dict:
        """
        Splits the text by section to create a dictionary with each section as the key to another dictionary holding the main topic, sub topic and text.
        """

        if self.pdf_text == None:
            return None

        # << Split the document into an array by "End of Document" >>
        pages = self._splitTextByEOD()

        # << Obtain a dictionary of the section, main heading, subheading and text >>
        # Main and Sub Heading
        main_heading = ''
        sub_heading = ''

        text_dict = {}

        for page in pages:

            # << Split the array to obtain the title and text data >>
            page_data = self._splitPageContentByBlackstone(page)

            # Only work on arrays with a length more than 1
            # At the very end of the document, there'll be a "End of Document" with a linespace after. This will register as the last element of the array
            if len(page_data) > 1:

                # << Get the page heading and text >>
                page_heading = self._getPageHeading(page_data)
                page_text = self._getPageText(page_data)

                # << Update Headings >>
                # Headings with all caps are main headings
                # Headings with standard letters are sub headings
                if self._isPageHeadingUpper(page_heading):
                    # Update Main Heading and Reset Sub Heading
                    main_heading = page_heading
                    sub_heading = ''
                else:
                    # Update Sub Heading
                    sub_heading = page_heading

                # << Get the sections and section text in the page text >>
                section_regex = re.escape(self.filename) + r'\.\d{1,3}\n'           # Create a regex to idenfity section headers (i.e. D4.52)
                sections = self._getPageSections(section_regex,page_text)           # Gets the list of sections that are in the page text
                section_texts = self._getPageSectionTexts(section_regex,page_text)  # Gets the list of section texts that are in the page text

                # Check if there are any sections on this page.
                # Only proceed if there are sections on the page.
                if len(sections) > 0:
                    # There are sections, proceed to split

                    for section, section_text in zip(sections,section_texts):

                        # << Format the section text into the desired format >>
                        section_text = self._formatSectionText(section_text,page_heading)

                        # << Get the Data in Dictionary Format >>
                        section_dict = self._getSectionDict(section,main_heading,sub_heading,section_text)

                        # << Update the main dictionary with the section dictionary >>
                        text_dict.update(section_dict)

        return text_dict

    def _getSectionDict(self,subsection:str,main_heading:str,sub_heading:str,subsection_text:List[str]) -> dict:
        """
        Create a dictionary with the data in the required format.

        PARAMETERS
        ----------

        subsection : str
            The subsection that will be used as a key for the dictionary.

        main_heading : str
            The main heading for this subsection.

        sub_heading : str
            The sub heading for this subsection.

        subsection_text : List[str]
            The array of text that is contained under this subsection.

        RETURNS
        -------
        Dictionary in the required format.

        dict {
            subsection : [
                "section_heading": main_heading,
                "section_subheading": sub_heading,
                "section_text": subsection_text
            ]
        }
        """
        return {
            subsection: {
                self.SECTION_MAIN_HEADING: main_heading,
                self.SECTION_SUB_HEADING: sub_heading,
                self.SECTION_TEXT: subsection_text
            }
        }

    def _formatSectionText(self,text:str,page_heading:str) -> List[str]:
        """
        Removes any occurrences of the page heading within the text and removes newline characters as deemed appropriate. For more information, look at the documentation for the function ProcedurePDF._removePageHeadingInText() and General.removeNewLine()

        PARAMETERS
        ----------
        text : str
            Raw text contained within the subsection.

        page_heading : str
            The heading that's on the current page of the PDF file.

        RETURNS
        -------
        A list of strings that represents the text separated by newline characters where necessary.

        List[str]

        SEE ALSO
        --------
        ProcedurePDF._removePageHeadingInText()
        ProcedurePDF.removeNewLine()
        """
        text = self._removePageHeadingInText(text,page_heading)
        text = self._removeNewLine(text)
        return text


    def _getPageHeading(self,page_data:List[str]) -> str:
        """
        Gets the page heading from an array containing the page data that has already been split with "Blackstone's Criminal Procedure 2022".
        """
        return page_data[0].replace('\n','')

    def _isPageHeadingUpper(self,heading:str) -> bool:
        """
        Checks if a large portion of the heading is in capital letters. This is because some headings that include sections will have a lower case s or numbers.

        If more than 80% of the characters are upper case, this will be considered as a main heading and return True.
        """
        alph = list(filter(str.isalpha,heading))
        percentage_uppercase = sum(map(str.isupper,alph)) / len(alph)

        if percentage_uppercase > 0.8: return True
        return False

    def _getPageText(self,page_data:List[str]) -> str:
        """
        Gets the page text from an array containing the page data that has already been split with "Blackstone's Criminal Procedure 2022".
        """       
        return page_data[1]

    def _getPageSections(self,regex:str,text:str) -> List[str]:
        """
        Gets the list of sections that are in the page text.
        """
        sections = re.findall(regex,text)                   # Gets all the section headers
        sections = [s.replace('\n','') for s in sections]   # Replace all the newline characters in the section headers
        return sections

    def _getPageSectionTexts(self,regex:str,text:str) -> List[str]:
        """
        Gets the list of section texts that are in the page text.
        """
        section_text = re.split(regex,text)     # Gets all the text for the section headers
        section_text.pop(0)                     # Removes the first item in the array - Text between the page title and first section on the page (Not Required)
        return section_text
    
    def _splitTextByEOD(self) -> List[str]:
        """
        Splits a string of text by the "End of Document" string into a List of strings.
        """
        eod_regex = r"End\sof\sDocument"
        eod_text = re.split(eod_regex,self.pdf_text)
        return eod_text

    def _splitPageContentByBlackstone(self,page:str) -> List[str]:
        """
        Takes in a string splits by "Blackstone's Criminal Practice 2022" as this occurs right after every main/sub heading on every page.
        """
        # Distinct Characteristics: Every heading/subheading will have "Blackstone's Criminal Practice 2022"
        # This will create an array of two elements - The title and the remainder of the text until end of document.
        # Title Regex    
        title_regex = r"\nBlackstone's\sCriminal\sPractice\s2022"
        page_data = re.split(title_regex,page)

        return page_data    

    def _removePageHeadingInText(self,text:str,heading:str) -> str:
        """
        Removes any occurances of the page heading followed by a newline character within the text. This function is implemented as if a document exceeds the length of the page, the page heading is repeated on the following page followed by a newline character. 
        
        When parsing through the PDF, the program will assume it is part of the text which is incorrect. This removes it.
        """
        return text.replace(f'{heading}\n','')

    def _removeNewLine(self,raw_text:str) -> List[str]:
        """
        Remove newline characters '\n' within the PDF text that are there due to space constraints. This avoids unnecessary line breaks in the middle of documents due to the formatting of the PDF document.
        
        This function will keep most linebreaks that occur at the end of paragraphs/bullet points.
        """

        linespace_regex = r'[\.;:—]\n|or\n|and\n'               # Regex to identify linespaces to be kept (New line after full stop etc.)
        delimiter_arr = re.findall(linespace_regex,raw_text)    # Getting an array of delimeters that should be kept
        text_arr = re.split(linespace_regex,raw_text)           # Getting an array of text to be rejoined to the delimiters
        
        text_arr_formatted = [s.replace('\n','') for s in text_arr]             # Replace '\n' with '' in the text array
        delimiter_arr_formatted = [delim.replace('\n','') for delim in delimiter_arr]   # Replace '\n' with '' in the delimiter array

        text_arr = []

        if len(delimiter_arr_formatted) > 0:
            # Delimiters present in text
            text_delim_zip = zip(text_arr_formatted,delimiter_arr_formatted)    # Zip both arrays together to form a dictionary
            
            # Using tuple unpacking, get the text and delimiter rejoined in an array where there should be a line break for every element within this array.
            for text, delimiter in text_delim_zip:
                text_arr.append(text + delimiter)
        else:
            # Delimiters not present in text
            text_arr += text_arr_formatted

        return text_arr
    
class Topic:

    """
    This class is built to handle each topic that is given and the respective sections that need to be referenced for this particular topic. One topic can reference from multiple files.
    """

    def __init__(self,topic:str,data:dict) -> None:
        self.topic = topic                              # Topic Number [str]
        self.data = data                                # Topic Data [dict]
        self.title = data['title']                      # Topic Title [str]
        self.sections_data = data['sections']           # Topic Sections and Subsections [dict]
        self.sections = self.sections_data.keys()       # Topic Sections [List[str]]

class DocxWriter:

    """
    This class handles the writing of the information to a word document.
    """

    def __init__(self,data:dict) -> None:
        self.doc = Document()
        self.doc_title = data['doc_title']  # Title of the document to be saved.
        self.doc_data = data['doc_data']    # Data dictionary of the topics and respective sections and subsections to extract.

        self.topics = []        # "List[Topic]"" - List of topic objects
        self.pdfs = {}          # "dict" - Dictionary of section and the PDFs (key: "[section]" as str | item: "[pdf_str]" as str)

    def createDocument(self) -> int:
        """
        Generates a document based on the data that is passed into the object.
        """

        # << Get Topics and PDFs >>
        if self._getTopicsAndPDFs() != 0:
            return self._getTopicsAndPDFs() # Error with generating topics and PDFs, refer to ._getTopicsAndPDFs()

        # << Write the Data for Topics >>
        for topic in self.topics:
            if self._writeTopicData(topic) != 0:
                return self._writeTopicData(topic) # Error with writing topic to word file, refer to ._writeTopicData()

            # Add a Page Break unless it's the last topic
            if topic != self.topics[-1]:
                self.doc.add_page_break()

        # Save the Document
        self.doc.save(f'output/{self.doc_title}.docx')

        # Return with Code 0 - Successful Generation of Document
        return 0

    def _writeTopicData(self,topic:Topic) -> int:
        logging.info(f'[Writing]: Writing data for {topic.topic}: {topic.title}')

        # Add the Topic Number and Topic Title as a "Title"
        self.doc.add_heading(f"{topic.topic}",level=0)
        self.doc.add_heading(f"{topic.title}",level=1)

        topic_data = {}

        # Iterate through the Sections and Subsections
        for section, subsections in topic.sections_data.items():
            logging.info(f"[Writing]: Getting PDF for {section}")
            pdf = self.pdfs[section]

            # Sort the Subsections in Ascending Order
            subsections.sort()

            section_dict = pdf.getSections(subsections)
            topic_data.update(section_dict)

        for section_title, section_data in topic_data.items():

            # Writing the Section as a Level 1 Heading
            section_main_heading = section_data[pdf.SECTION_MAIN_HEADING]
            section_sub_heading = section_data[pdf.SECTION_SUB_HEADING]
            section_text = section_data[pdf.SECTION_TEXT]

            section_heading = f"{section_title} - {section_main_heading}"

            if section_sub_heading != '':
                # Sub Heading is not Empty
                if section_main_heading == "":
                    # Main Heading is Empty
                    section_heading = f"{section_title} - {section_sub_heading}"
                else:
                    section_heading += f" > {section_sub_heading}"

            logging.info(f"[Writing]: Writing for Subsection {section_title}")
            self.doc.add_heading(section_heading,level=2)

            # Writing the text data to the document
            for text_item in section_text:
                para = self.doc.add_paragraph(text_item)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                para.paragraph_format.space_after = Pt(6)

        return 0

    def _getTopicsAndPDFs(self) -> int:
        """
        Fills the pdfs and topics property of this class.
        """
        sections = []

        for topic_name, topic_data in self.doc_data.items():
            # key: Topic # [str]
            # items: Topic Data [dict]
            
            topic = Topic(topic_name,topic_data)    # Create Topic Object
            self.topics.append(topic)               # Append Topic Object to Topic Array

            sections += topic.sections              # Get all the sections for the topic

        unique_sections = list(set(sections))       # Convert the list for all the sections into a unique list so no repetitions

        if self._getPDFObjects(unique_sections) != 0:               # Update self.pdfs 
            return 1                                                # Return 1 for unsuccessful PDF conversion

        return 0
    
    def _getPDFObjects(self,unique_sections:List[str]) -> int:

        for section in unique_sections:
            logging.info(f'[PDFs]: Loading PDF for {section}')
            current_pdf = {section: ProcedurePDF(section)}
            self.pdfs.update(current_pdf)

        logging.info("<< ALL PDF FILES LOADED >>")

        return 0

if __name__ == '__main__':

    # Handling the Logger
    log = Logger()

    with open('topics.json') as f:
        log.logInfo('Loading JSON Data.')
        data = json.load(f)
    
    writer = DocxWriter(data)
    writer.createDocument()